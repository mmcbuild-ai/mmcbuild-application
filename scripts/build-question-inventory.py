"""
Cross-module question inventory generator.

Produces:
  - briefs/outputs/mmcbuild_question_inventory_v1.xlsx
  - briefs/outputs/mmcbuild_question_inventory_v1.docx

The question records below are extracted verbatim from the source files at the
SHA captured in SOURCE_COMMIT. Re-run after edits to regenerate.

Run:
  python scripts/build-question-inventory.py
"""

from __future__ import annotations

import datetime as dt
import subprocess
from pathlib import Path

from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.formatting.rule import FormulaRule
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.table import Table, TableStyleInfo

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


REPO_ROOT = Path(__file__).resolve().parent.parent
OUT_DIR = REPO_ROOT / "briefs" / "outputs"
XLSX_PATH = OUT_DIR / "mmcbuild_question_inventory_v1.xlsx"
DOCX_PATH = OUT_DIR / "mmcbuild_question_inventory_v1.docx"

HEADER_FILL_HEX = "1F3A5F"
HEADER_FONT_HEX = "FFFFFF"


def get_commit_sha() -> str:
    try:
        out = subprocess.check_output(
            ["git", "rev-parse", "HEAD"], cwd=REPO_ROOT, stderr=subprocess.DEVNULL
        )
        return out.decode().strip()
    except Exception:
        return "unknown"


def Q(
    module: str,
    section: str,
    order: int,
    qid: str,
    label: str,
    qtype: str,
    source_file: str,
    source_line: int,
    helper_text: str = "",
    options: str = "",
    required: bool = False,
    conditional_on: str = "",
) -> dict:
    return {
        "module": module,
        "section": section,
        "order": order,
        "id": qid,
        "label": label,
        "helper_text": helper_text,
        "type": qtype,
        "options": options,
        "required": required,
        "conditional_on": conditional_on,
        "source_file": source_file,
        "source_line": source_line,
    }


# ---- Option lists (verbatim from source) -----------------------------------

DESIGN_STAGES = "Concept / brief; Schematic design; Design development; Documentation (DA-ready); Construction Certificate ready; Submitted (post-DA); Submitted (post-CC)"
PROJECT_GOALS_OPTS = "Explore MMC options early; Validate compliance pathway; Compare cost vs. traditional; Brief the client; Submission-ready evidence pack; Educate myself on MMC"
SUBMISSION_TIMELINES = "No fixed date; Within 4 weeks; 1–3 months; 3–6 months; Already submitted"
BUILDING_TYPOLOGIES = "Single residential; Duplex; Townhouse; Apartment; Hotel; Mixed use; Commercial"
BUILDING_CLASSES = "Class 1a; Class 1b; Class 10a; Class 10b"
CONSTRUCTION_TYPES = "Type A; Type B; Type C"
IMPORTANCE_LEVELS = "1; 2; 3; 4"
SOIL_CLASSIFICATIONS = "A; S; M; M-D; H1; H2; E; P"
FOOTING_TYPES = "Strip footing; Pad footing; Raft slab; Waffle slab; Stiffened raft; Stumps/Piers; Screw piles"
FRAMING_MATERIALS = "Timber; Steel; Timber + Steel hybrid; Masonry; SIPs"
WIND_CLASSIFICATIONS = "N1; N2; N3; N4; N5; N6; C1; C2; C3; C4"
TERRAIN_CATEGORIES = "TC1; TC2; TC2.5; TC3"
ROOF_MATERIALS = "Concrete tile; Terracotta tile; Metal (Colorbond); Metal (Zincalume); Slate; Asphalt shingle"
WALL_CLADDINGS = "Brick veneer; Double brick; Fibre cement; Timber weatherboard; Metal cladding; Rendered foam; Autoclaved aerated concrete"
DPC_TYPES = "Polyethylene membrane; Bituminous membrane; Chemical DPC; Not specified"
GARAGE_LOCATIONS = "Attached; Detached; Integrated/under main roof; N/A"
SMOKE_ALARM_TYPES = "Photoelectric (hardwired interconnected); Photoelectric (battery); Ionisation; Combined photo/ion"
VENTILATION_METHODS = "Openable windows; Openable windows + ceiling fans; Mechanical ventilation; Mixed mode"
ENERGY_PATHWAYS = "DTS (Deemed-to-Satisfy); NatHERS; JV3 (Verification)"
GLAZING_TYPES = "Single clear; Single tinted; Double glazed (clear); Double glazed (low-e); Triple glazed"
HOT_WATER_SYSTEMS = "Electric storage; Electric heat pump; Gas storage; Gas instantaneous; Solar electric boost; Solar gas boost"
HEATING_TYPES = "Ducted gas; Ducted reverse cycle; Split system; Hydronic; Wood heater (open flue); Wood heater (closed flue); Electric panel"
CLIMATE_ZONES = "1; 2; 3; 4; 5; 6; 7; 8"
BAL_RATINGS = "N/A; BAL-LOW; BAL-12.5; BAL-19; BAL-29; BAL-40; BAL-FZ"
PROJECT_STATUSES = "Draft; Active; Completed; Archived"
DISCIPLINES = "Architect; Structural Engineer; Hydraulic Engineer; Energy Consultant; Building Surveyor; Geotechnical Engineer; Acoustic Engineer; Fire Engineer; Landscape Architect; Builder; Other"
CONSTRUCTION_SYSTEMS_OPTS = "SIPs; CLT / Mass Timber; Steel Frame; Timber Frame; Volumetric Modular; Hybrid"
QUOTE_REGIONS = "NSW; VIC; QLD; WA; SA; TAS; ACT; NT"
RESPONSE_STATUSES = "acknowledged (I have reviewed this finding); in_progress (Remediation work has started); completed (Remediation is complete); disputed (I disagree with this finding)"
THUMB_RATING = "Yes (1); No (-1)"
THREE_RATING = "Correct (1); Partially correct (0); Incorrect (-1)"

PROJECT_FORM = "src/components/projects/questionnaire-form.tsx"
CREATE_DLG = "src/components/projects/create-project-dialog.tsx"
EDIT_DLG = "src/components/projects/edit-project-dialog.tsx"
CONTRIB_FORM = "src/components/projects/project-contributors.tsx"
FEEDBACK_WIDGET = "src/components/comply/feedback-widget.tsx"
FINDING_FEEDBACK = "src/components/comply/finding-feedback.tsx"
AMEND_DLG = "src/components/comply/finding-amend-dialog.tsx"
REJECT_DLG = "src/components/comply/finding-reject-dialog.tsx"
SHARE_DLG = "src/components/comply/share-finding-dialog.tsx"
RESPOND_FORM = "src/app/respond/[token]/response-form.tsx"
SYS_PANEL = "src/components/build/system-selection-panel.tsx"
ESTIMATE_BTN = "src/components/quote/run-estimate-button.tsx"
HOLDING_CALC = "src/components/quote/holding-cost-calculator.tsx"


# ---- Question records ------------------------------------------------------

QUESTIONS: list[dict] = [
    # PROJECT — Create dialog
    Q("Project", "Project Creation", 1, "name", "Project Name", "free_text", CREATE_DLG, 101,
      helper_text="e.g. 42 Smith Street Renovation", required=True),
    Q("Project", "Project Creation", 2, "address", "Address", "free_text", CREATE_DLG, 110,
      helper_text="Australian address; auto-derives climate zone, wind region, council, BAL via Mapbox + property services",
      required=True),

    # PROJECT — Edit dialog (name/address overlap with Create — captured once; status is edit-only)
    Q("Project", "Project Edit", 1, "status", "Project Status", "single_select", EDIT_DLG, 33,
      options=PROJECT_STATUSES, required=True,
      helper_text="Workflow status of the project record (set on edit, not creation)"),

    # PROJECT — Onboarding questionnaire, Step 0: Project Status
    Q("Project", "Project Status", 1, "design_stage", "What stage are your designs at?", "single_select",
      PROJECT_FORM, 421, options=DESIGN_STAGES,
      helper_text="MMC analysis pays off most at concept and schematic stages — before drawings lock and before council submission."),
    Q("Project", "Project Status", 2, "project_goals", "What do you want to get out of this project?", "multi_select",
      PROJECT_FORM, 441, options=PROJECT_GOALS_OPTS,
      helper_text="Select all that apply. Stored pipe-separated."),
    Q("Project", "Project Status", 3, "submission_timeline", "Submission timeline", "single_select",
      PROJECT_FORM, 474, options=SUBMISSION_TIMELINES,
      helper_text="When are you targeting council submission?"),

    # PROJECT — Step 1: Building Classification
    Q("Project", "Building Classification", 1, "building_typology", "Building Typology", "single_select",
      PROJECT_FORM, 487, options=BUILDING_TYPOLOGIES),
    Q("Project", "Building Classification", 2, "building_class", "Building Classification (NCC)", "single_select",
      PROJECT_FORM, 493, options=BUILDING_CLASSES),
    Q("Project", "Building Classification", 3, "construction_type", "Construction Type", "single_select",
      PROJECT_FORM, 499, options=CONSTRUCTION_TYPES),
    Q("Project", "Building Classification", 4, "importance_level", "Importance Level", "single_select",
      PROJECT_FORM, 505, options=IMPORTANCE_LEVELS),

    # PROJECT — Step 2: Structure & Footings (H1)
    Q("Project", "Structure & Footings (H1)", 1, "storeys", "Number of Storeys", "number",
      PROJECT_FORM, 517, helper_text="Min 1, max 10"),
    Q("Project", "Structure & Footings (H1)", 2, "floor_area", "Total Floor Area (m²)", "number",
      PROJECT_FORM, 523),
    Q("Project", "Structure & Footings (H1)", 3, "soil_classification", "Soil Classification (AS 2870)", "single_select",
      PROJECT_FORM, 532, options=SOIL_CLASSIFICATIONS, helper_text="To be confirmed"),
    Q("Project", "Structure & Footings (H1)", 4, "footing_type", "Expected Footing Type", "single_select",
      PROJECT_FORM, 539, options=FOOTING_TYPES,
      helper_text="Select if known — MMC product selection may change this"),
    Q("Project", "Structure & Footings (H1)", 5, "framing_material", "Framing Material", "single_select",
      PROJECT_FORM, 546, options=FRAMING_MATERIALS),
    Q("Project", "Structure & Footings (H1)", 6, "wind_classification", "Wind Classification (AS 4055)", "single_select",
      PROJECT_FORM, 552, options=WIND_CLASSIFICATIONS,
      helper_text="Auto-derived from address if not overridden"),
    Q("Project", "Structure & Footings (H1)", 7, "terrain_category", "Terrain Category", "single_select",
      PROJECT_FORM, 559, options=TERRAIN_CATEGORIES),

    # PROJECT — Step 3: Weatherproofing (H2)
    Q("Project", "Weatherproofing (H2)", 1, "roof_material", "Roof Material", "single_select",
      PROJECT_FORM, 570, options=ROOF_MATERIALS),
    Q("Project", "Weatherproofing (H2)", 2, "wall_cladding", "Wall Cladding", "single_select",
      PROJECT_FORM, 576, options=WALL_CLADDINGS),
    Q("Project", "Weatherproofing (H2)", 3, "dpc_type", "Damp-Proof Course / Waterproofing Membrane", "single_select",
      PROJECT_FORM, 582, options=DPC_TYPES),
    Q("Project", "Weatherproofing (H2)", 4, "sarking", "Roof sarking installed (optional)", "boolean",
      PROJECT_FORM, 588),
    Q("Project", "Weatherproofing (H2)", 5, "subfloor_ventilation", "Sub-floor ventilation (optional)", "boolean",
      PROJECT_FORM, 593),

    # PROJECT — Step 4: Fire Safety (H3)
    Q("Project", "Fire Safety (H3)", 1, "distance_to_boundary", "Distance to Boundary (m)", "number",
      PROJECT_FORM, 603, helper_text="e.g., 1.5"),
    Q("Project", "Fire Safety (H3)", 2, "attached_dwelling", "Attached dwelling (party wall)", "boolean",
      PROJECT_FORM, 612,
      conditional_on="building_typology in {Duplex, Townhouse, Apartment, Mixed use}"),
    Q("Project", "Fire Safety (H3)", 3, "garage_location", "Garage Location", "single_select",
      PROJECT_FORM, 618, options=GARAGE_LOCATIONS),
    Q("Project", "Fire Safety (H3)", 4, "smoke_alarm_type", "Smoke Alarm Type", "single_select",
      PROJECT_FORM, 624, options=SMOKE_ALARM_TYPES),
    Q("Project", "Fire Safety (H3)", 5, "party_wall_frl", "Party Wall FRL (e.g., 60/60/60)", "free_text",
      PROJECT_FORM, 631,
      conditional_on="attached_dwelling=true AND multi-dwelling typology",
      helper_text="e.g., 60/60/60"),

    # PROJECT — Step 5: Health & Amenity (H4)
    Q("Project", "Health & Amenity (H4)", 1, "wet_area_count", "Number of Wet Areas", "number",
      PROJECT_FORM, 643),
    Q("Project", "Health & Amenity (H4)", 2, "ceiling_height_habitable", "Ceiling Height — Habitable Rooms (m)", "number",
      PROJECT_FORM, 650, helper_text="Default 2.4; min 2.1, max 4"),
    Q("Project", "Health & Amenity (H4)", 3, "ceiling_height_non_habitable", "Ceiling Height — Non-habitable Rooms (m)", "number",
      PROJECT_FORM, 659, helper_text="Default 2.1; min 2.1, max 4"),
    Q("Project", "Health & Amenity (H4)", 4, "exhaust_fans", "Exhaust fans to all wet areas", "boolean",
      PROJECT_FORM, 668),
    Q("Project", "Health & Amenity (H4)", 5, "natural_ventilation_method", "Natural Ventilation Method", "single_select",
      PROJECT_FORM, 673, options=VENTILATION_METHODS),

    # PROJECT — Step 6: Energy Efficiency (H6)
    Q("Project", "Energy Efficiency (H6)", 1, "energy_pathway", "Energy Compliance Pathway", "single_select",
      PROJECT_FORM, 684, options=ENERGY_PATHWAYS),
    Q("Project", "Energy Efficiency (H6)", 2, "insulation_ceiling_r", "Ceiling Insulation R-value", "number",
      PROJECT_FORM, 690, helper_text="e.g., 6.0"),
    Q("Project", "Energy Efficiency (H6)", 3, "insulation_wall_r", "Wall Insulation R-value", "number",
      PROJECT_FORM, 698, helper_text="e.g., 2.5"),
    Q("Project", "Energy Efficiency (H6)", 4, "insulation_floor_r", "Floor Insulation R-value", "number",
      PROJECT_FORM, 706, helper_text="e.g., 1.0 (0 if slab on ground)"),
    Q("Project", "Energy Efficiency (H6)", 5, "glazing_type", "Glazing Type", "single_select",
      PROJECT_FORM, 714, options=GLAZING_TYPES),
    Q("Project", "Energy Efficiency (H6)", 6, "hot_water_system", "Hot Water System", "single_select",
      PROJECT_FORM, 720, options=HOT_WATER_SYSTEMS),
    Q("Project", "Energy Efficiency (H6)", 7, "has_solar_pv", "Solar PV installed", "boolean",
      PROJECT_FORM, 726),
    Q("Project", "Energy Efficiency (H6)", 8, "nathers_rating", "NatHERS Star Rating", "number",
      PROJECT_FORM, 732, conditional_on="energy_pathway=NatHERS",
      helper_text="0–10, e.g., 7.0"),

    # PROJECT — Step 7: Site, Climate & Bushfire
    Q("Project", "Site, Climate & Bushfire", 1, "climate_zone", "Climate Zone", "single_select",
      PROJECT_FORM, 747, options=CLIMATE_ZONES,
      helper_text="Auto-derived from address if not overridden"),
    Q("Project", "Site, Climate & Bushfire", 2, "bal_rating", "Bushfire Attack Level (BAL)", "single_select",
      PROJECT_FORM, 754, options=BAL_RATINGS,
      helper_text="Auto-derived from address if not overridden"),
    Q("Project", "Site, Climate & Bushfire", 3, "site_conditions", "Site Conditions", "free_text",
      PROJECT_FORM, 760,
      helper_text="Used to flag flood overlay, slope, and setback constraints during compliance and 3D layout. e.g., flat site, no flood overlay, corner block"),
    Q("Project", "Site, Climate & Bushfire", 4, "has_swimming_pool", "Swimming pool on site", "boolean",
      PROJECT_FORM, 768),
    Q("Project", "Site, Climate & Bushfire", 5, "has_heating_appliance", "Heating appliance (wood heater, gas fire, etc.)", "boolean",
      PROJECT_FORM, 773),
    Q("Project", "Site, Climate & Bushfire", 6, "heating_type", "Heating Type", "single_select",
      PROJECT_FORM, 779, options=HEATING_TYPES,
      conditional_on="has_heating_appliance=true"),

    # PROJECT — Step 8: Access & Livable Housing (H5/H8) — conditional on residential typology
    Q("Project", "Access & Livable Housing (H5/H8)", 1, "has_stairs", "Has stairs", "boolean",
      PROJECT_FORM, 791,
      conditional_on="building_typology in {Single residential, Duplex, Townhouse, Apartment, Mixed use}"),
    Q("Project", "Access & Livable Housing (H5/H8)", 2, "has_balcony_deck", "Has balcony or deck", "boolean",
      PROJECT_FORM, 796,
      conditional_on="building_typology in {Single residential, Duplex, Townhouse, Apartment, Mixed use}"),
    Q("Project", "Access & Livable Housing (H5/H8)", 3, "max_fall_height", "Maximum Fall Height (m)", "number",
      PROJECT_FORM, 802,
      conditional_on="has_balcony_deck=true",
      helper_text="e.g., 1.0"),
    Q("Project", "Access & Livable Housing (H5/H8)", 4, "has_step_free_entry", "Step-free entry provided", "boolean",
      PROJECT_FORM, 811),
    Q("Project", "Access & Livable Housing (H5/H8)", 5, "accessible_bathroom", "Accessible bathroom (Livable Housing)", "boolean",
      PROJECT_FORM, 816),
    Q("Project", "Access & Livable Housing (H5/H8)", 6, "min_door_width", "Minimum Door Width (mm)", "number",
      PROJECT_FORM, 821, helper_text="e.g., 820"),
    Q("Project", "Access & Livable Housing (H5/H8)", 7, "min_corridor_width", "Minimum Corridor Width (mm)", "number",
      PROJECT_FORM, 829, helper_text="e.g., 1000"),
    Q("Project", "Access & Livable Housing (H5/H8)", 8, "services", "Services", "free_text",
      PROJECT_FORM, 837,
      helper_text="Used to plan service routing in the 3D model and trigger service-specific compliance checks. e.g., smoke alarms, ducted A/C, gas cooktop"),
    Q("Project", "Access & Livable Housing (H5/H8)", 9, "special_requirements", "Special Requirements", "free_text",
      PROJECT_FORM, 844,
      helper_text="Heritage controls, accessibility upgrades, and other constraints that should shape optimisation outputs."),

    # PROJECT — Project Team (Add Contributor)
    Q("Project", "Project Team (Contributors)", 1, "contact_name", "Contact Name", "free_text",
      CONTRIB_FORM, 229, required=True),
    Q("Project", "Project Team (Contributors)", 2, "discipline", "Discipline", "single_select",
      CONTRIB_FORM, 239, options=DISCIPLINES, required=True),
    Q("Project", "Project Team (Contributors)", 3, "company_name", "Company", "free_text",
      CONTRIB_FORM, 258),
    Q("Project", "Project Team (Contributors)", 4, "contact_email", "Email", "free_text",
      CONTRIB_FORM, 268),
    Q("Project", "Project Team (Contributors)", 5, "contact_phone", "Phone", "free_text",
      CONTRIB_FORM, 277),
    Q("Project", "Project Team (Contributors)", 6, "notes", "Notes", "free_text",
      CONTRIB_FORM, 288),

    # COMPLY — note: the Project questionnaire is re-exported by Comply; recorded once under Project.

    # COMPLY — Workflow inputs (operational forms; flagged for Karen to consider separately from intake)
    Q("Comply", "Workflow — Report Feedback", 1, "report_helpful", "Was this report helpful?", "single_select",
      FEEDBACK_WIDGET, 39, options=THUMB_RATING),
    Q("Comply", "Workflow — Report Feedback", 2, "report_comment", "Optional comment", "free_text",
      FEEDBACK_WIDGET, 67),

    Q("Comply", "Workflow — Finding Feedback", 1, "finding_accurate", "Accurate?", "single_select",
      FINDING_FEEDBACK, 40, options=THREE_RATING),

    Q("Comply", "Workflow — Finding Amend", 1, "amended_description", "Amended Description", "free_text",
      AMEND_DLG, 91, helper_text="Original wording is preserved alongside the amendment."),
    Q("Comply", "Workflow — Finding Amend", 2, "amended_action", "Amended Remediation Action", "free_text",
      AMEND_DLG, 101),
    Q("Comply", "Workflow — Finding Amend", 3, "amended_discipline", "Discipline", "single_select",
      AMEND_DLG, 111, options=DISCIPLINES),
    Q("Comply", "Workflow — Finding Amend", 4, "assigned_contributor_id", "Assign Contributor", "single_select",
      AMEND_DLG, 131,
      helper_text="From project's contributor list",
      conditional_on="project has ≥1 contributor"),

    Q("Comply", "Workflow — Finding Reject", 1, "rejection_reason", "Rejection Reason", "free_text",
      REJECT_DLG, 66, required=True,
      helper_text="Explain why this finding is being rejected. Recorded in the activity log."),

    Q("Comply", "Workflow — Finding Share (new contributor)", 1, "share_name", "Name", "free_text",
      SHARE_DLG, 169, required=True),
    Q("Comply", "Workflow — Finding Share (new contributor)", 2, "share_email", "Email", "free_text",
      SHARE_DLG, 179, required=True),
    Q("Comply", "Workflow — Finding Share (new contributor)", 3, "share_company", "Company", "free_text",
      SHARE_DLG, 190),
    Q("Comply", "Workflow — Finding Share (new contributor)", 4, "share_discipline", "Discipline", "single_select",
      SHARE_DLG, 200, options=DISCIPLINES),

    Q("Comply", "Contributor Remediation Response (token-based)", 1, "remediation_status", "Remediation Status", "single_select",
      RESPOND_FORM, 178, options=RESPONSE_STATUSES, required=True,
      helper_text="External-facing form at /respond/[token]; not behind Supabase auth."),
    Q("Comply", "Contributor Remediation Response (token-based)", 2, "remediation_notes", "Notes", "free_text",
      RESPOND_FORM, 210, helper_text="Describe the remediation work done or provide additional context"),
    Q("Comply", "Contributor Remediation Response (token-based)", 3, "remediation_file", "Supporting Document", "file_upload",
      RESPOND_FORM, 224,
      helper_text="Optional. Max 10MB. PDF, DOC, images, or spreadsheets."),

    # BUILD
    Q("Build", "Construction Systems", 1, "selected_systems", "Construction Systems (select one or more)", "multi_select",
      SYS_PANEL, 11, options=CONSTRUCTION_SYSTEMS_OPTS,
      helper_text="Flows into Comply, Quote, Directory, and Training. Stored on the project record."),

    # QUOTE
    Q("Quote", "Pre-flight (Cost Estimation)", 1, "region", "Project Region", "single_select",
      ESTIMATE_BTN, 21, options=QUOTE_REGIONS,
      helper_text="Selected at run-time when starting an estimate; defaults to NSW."),

    Q("Quote", "Holding Cost Calculator", 1, "weekly_finance_cost", "Finance / Interest", "number",
      HOLDING_CALC, 211, helper_text="$/week"),
    Q("Quote", "Holding Cost Calculator", 2, "weekly_site_costs", "Site Running Costs", "number",
      HOLDING_CALC, 216, helper_text="$/week"),
    Q("Quote", "Holding Cost Calculator", 3, "weekly_insurance", "Insurance", "number",
      HOLDING_CALC, 221, helper_text="$/week"),
    Q("Quote", "Holding Cost Calculator", 4, "weekly_opportunity_cost", "Opportunity Cost (Lost Rent)", "number",
      HOLDING_CALC, 226, helper_text="$/week"),
    Q("Quote", "Holding Cost Calculator", 5, "weekly_council_fees", "Council & Permit Fees", "number",
      HOLDING_CALC, 231, helper_text="$/week"),
    Q("Quote", "Holding Cost Calculator", 6, "custom_items", "Custom Items (label + $/week)", "free_text",
      HOLDING_CALC, 244,
      helper_text="Repeatable rows, e.g. \"Storage rental\". Stored as JSON array of {label, amount}."),
]


GAPS: list[str] = [
    "Comply has no module-specific intake questions of its own — it consumes the Project questionnaire (verified: src/components/comply/questionnaire-form.tsx is a re-export of src/components/projects/questionnaire-form.tsx).",
    "Build has no project-questionnaire-style intake — only the Construction Systems multi-select. Karen may want to flag whether system selection belongs in Project intake (so Comply/Quote both see it) or remain in Build.",
    "Quote's only intake is the Region selector at run-time and the Holding Cost Calculator on the report page. Region duplicates what could be derived from the Project address; consider unifying.",
    "Project Status / design_stage / submission_timeline (Step 0) overlap conceptually with Project record's `status` field (set in Edit dialog). Two sources of truth for project lifecycle.",
    "Wind classification, climate zone, and BAL are auto-derived from the address but still presented as questions; Karen may decide whether to surface them as confirmable values vs. hidden auto-fields.",
    "Workflow inputs (feedback widgets, finding amend/reject/share dialogs, contributor response form) are operational forms, not user-facing intake — listed under Comply for completeness; Karen may want to filter these out when reordering intake questions.",
]


# ---- Excel generation ------------------------------------------------------

COLUMNS = [
    ("module", "Module", 14),
    ("section", "Section", 36),
    ("order", "Order", 7),
    ("id", "ID", 28),
    ("label", "Label", 50),
    ("helper_text", "Helper Text", 50),
    ("type", "Type", 14),
    ("options", "Options", 50),
    ("required", "Required", 10),
    ("conditional_on", "Conditional On", 40),
    ("source_file", "Source File", 56),
    ("source_line", "Source Line", 11),
    # Karen working columns:
    ("verdict", "Karen's Verdict", 18),
    ("new_module", "Suggested New Module", 22),
    ("new_section", "Suggested New Section", 28),
    ("notes", "Notes", 40),
]


def _style_header(ws, ncols: int) -> None:
    fill = PatternFill("solid", fgColor=HEADER_FILL_HEX)
    font = Font(bold=True, color=HEADER_FONT_HEX)
    for c in range(1, ncols + 1):
        cell = ws.cell(row=1, column=c)
        cell.fill = fill
        cell.font = font
        cell.alignment = Alignment(vertical="center", wrap_text=True)


def _write_question_sheet(ws, rows: list[dict], include_working_cols: bool) -> None:
    cols = COLUMNS if include_working_cols else COLUMNS[:12]
    headers = [h for _, h, _ in cols]
    ws.append(headers)
    for q in rows:
        ws.append([
            q.get("module", ""),
            q.get("section", ""),
            q.get("order", ""),
            q.get("id", ""),
            q.get("label", ""),
            q.get("helper_text", ""),
            q.get("type", ""),
            q.get("options", ""),
            "TRUE" if q.get("required") else "FALSE",
            q.get("conditional_on", ""),
            q.get("source_file", ""),
            q.get("source_line", ""),
        ] + ([""] * 4 if include_working_cols else []))

    for idx, (_, _, width) in enumerate(cols, start=1):
        ws.column_dimensions[get_column_letter(idx)].width = width

    _style_header(ws, len(cols))
    ws.freeze_panes = "B2"
    last_col = get_column_letter(len(cols))
    last_row = len(rows) + 1
    ws.auto_filter.ref = f"A1:{last_col}{last_row}"

    # Wrap text body
    for row in ws.iter_rows(min_row=2, max_row=last_row, max_col=len(cols)):
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)


def build_excel() -> None:
    wb = Workbook()
    summary = wb.active
    summary.title = "Summary"

    summary.append(["MMC Build — Cross-Module Question Inventory v1"])
    summary["A1"].font = Font(bold=True, size=14, color=HEADER_FILL_HEX)
    summary.append([])
    summary.append(["Generated", dt.datetime.now().isoformat(timespec="seconds")])
    summary.append(["Source commit", get_commit_sha()])
    summary.append(["Generated by", "Claude Code (instructed by Dennis McMahon)"])
    summary.append([])

    summary.append(["Module", "Question Count", "Sections"])
    by_module: dict[str, dict] = {}
    for q in QUESTIONS:
        m = q["module"]
        by_module.setdefault(m, {"count": 0, "sections": set()})
        by_module[m]["count"] += 1
        by_module[m]["sections"].add(q["section"])

    total = 0
    total_sections = set()
    for m in ["Project", "Comply", "Build", "Quote"]:
        info = by_module.get(m, {"count": 0, "sections": set()})
        summary.append([m, info["count"], len(info["sections"])])
        total += info["count"]
        total_sections |= {(m, s) for s in info["sections"]}
    summary.append(["TOTAL", total, len(total_sections)])

    # Bold the table header row
    header_row = 7
    for c in range(1, 4):
        summary.cell(row=header_row, column=c).font = Font(bold=True, color=HEADER_FONT_HEX)
        summary.cell(row=header_row, column=c).fill = PatternFill("solid", fgColor=HEADER_FILL_HEX)

    summary.append([])
    summary.append(["Gaps detected"])
    summary.cell(row=summary.max_row, column=1).font = Font(bold=True)
    for g in GAPS:
        summary.append(["", g])

    for col, width in [("A", 22), ("B", 90), ("C", 12)]:
        summary.column_dimensions[col].width = width
    for row in summary.iter_rows():
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)

    # All Questions
    ws_all = wb.create_sheet("All Questions")
    _write_question_sheet(ws_all, QUESTIONS, include_working_cols=True)

    # Per-module
    for m in ["Project", "Comply", "Build", "Quote"]:
        ws_m = wb.create_sheet(m)
        rows = [q for q in QUESTIONS if q["module"] == m]
        _write_question_sheet(ws_m, rows, include_working_cols=True)

    # Re-ordering Workspace = copy of All Questions with conditional formatting
    ws_re = wb.create_sheet("Re-ordering Workspace")
    _write_question_sheet(ws_re, QUESTIONS, include_working_cols=True)
    last_row = len(QUESTIONS) + 1
    ncols = len(COLUMNS)
    last_col = get_column_letter(ncols)
    yellow = PatternFill("solid", fgColor="FFF3CD")
    green = PatternFill("solid", fgColor="D4EDDA")
    red = PatternFill("solid", fgColor="F8D7DA")

    rng = f"A2:{last_col}{last_row}"
    ws_re.conditional_formatting.add(
        rng,
        FormulaRule(formula=['AND($N2<>"",$N2<>$A2)'], fill=yellow),
    )
    ws_re.conditional_formatting.add(
        rng,
        FormulaRule(formula=['$M2="Keep"'], fill=green),
    )
    ws_re.conditional_formatting.add(
        rng,
        FormulaRule(formula=['$M2="Remove"'], fill=red),
    )

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    wb.save(XLSX_PATH)


# ---- Word generation -------------------------------------------------------

def _set_heading_color(paragraph, hex_color: str) -> None:
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor.from_string(hex_color)


def build_word() -> None:
    doc = Document()

    # Page setup: A4, 1-inch margins
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("MMC Build — Cross-Module Question Inventory", level=0)
    _set_heading_color(title, HEADER_FILL_HEX)

    sub = doc.add_paragraph()
    sub_run = sub.add_run("Working document for Karen Burns review")
    sub_run.italic = True

    meta = doc.add_paragraph()
    meta.add_run(f"Generated: {dt.date.today().isoformat()}\n").bold = False
    meta.add_run(f"Source commit: {get_commit_sha()}\n")
    meta.add_run("Companion to: mmcbuild_question_inventory_v1.xlsx")

    # Purpose
    h = doc.add_heading("1. Purpose", level=1)
    _set_heading_color(h, HEADER_FILL_HEX)
    doc.add_paragraph(
        "This document is a read-only snapshot of every user-facing question across the "
        "Project (pre-module onboarding), Comply, Build, and Quote modules. Use it as the "
        "reference companion to the paired Excel workbook, where re-ordering decisions are "
        "captured. The Excel sheet \"Re-ordering Workspace\" has working columns "
        "(Karen's Verdict, Suggested New Module, Suggested New Section, Notes) for marking "
        "moves; this Word doc is for context."
    )

    # Summary
    h = doc.add_heading("2. Summary", level=1)
    _set_heading_color(h, HEADER_FILL_HEX)

    summary_table = doc.add_table(rows=1, cols=3)
    summary_table.style = "Light Grid Accent 1"
    hdr = summary_table.rows[0].cells
    hdr[0].text = "Module"
    hdr[1].text = "Question Count"
    hdr[2].text = "Sections"

    by_module: dict[str, dict] = {}
    for q in QUESTIONS:
        m = q["module"]
        by_module.setdefault(m, {"count": 0, "sections": set()})
        by_module[m]["count"] += 1
        by_module[m]["sections"].add(q["section"])

    total = 0
    total_sections = set()
    for m in ["Project", "Comply", "Build", "Quote"]:
        info = by_module.get(m, {"count": 0, "sections": set()})
        row = summary_table.add_row().cells
        row[0].text = m
        row[1].text = str(info["count"])
        row[2].text = str(len(info["sections"]))
        total += info["count"]
        total_sections |= {(m, s) for s in info["sections"]}
    row = summary_table.add_row().cells
    row[0].text = "TOTAL"
    row[1].text = str(total)
    row[2].text = str(len(total_sections))

    # Per module
    section_index = 3
    for m in ["Project", "Comply", "Build", "Quote"]:
        h = doc.add_heading(f"{section_index}. {m} module", level=1)
        _set_heading_color(h, HEADER_FILL_HEX)
        section_index += 1

        rows = [q for q in QUESTIONS if q["module"] == m]
        if not rows:
            doc.add_paragraph(
                f"No questions found in {m} module. See Gaps and observations."
            )
            continue

        # Group by section preserving first-occurrence order
        seen_sections: list[str] = []
        for q in rows:
            if q["section"] not in seen_sections:
                seen_sections.append(q["section"])

        for sec_idx, sec_name in enumerate(seen_sections, start=1):
            sh = doc.add_heading(f"{section_index - 1}.{sec_idx} {sec_name}", level=2)
            _set_heading_color(sh, HEADER_FILL_HEX)
            sec_questions = [q for q in rows if q["section"] == sec_name]
            sec_questions.sort(key=lambda x: x["order"])
            for q in sec_questions:
                p = doc.add_paragraph()
                run = p.add_run(q["label"])
                run.bold = True
                if q.get("required"):
                    p.add_run(" *").bold = True

                meta_lines = [
                    f"id: {q['id']}",
                    f"type: {q['type']}",
                ]
                if q.get("options"):
                    meta_lines.append(f"options: {q['options']}")
                if q.get("conditional_on"):
                    meta_lines.append(f"conditional_on: {q['conditional_on']}")
                if q.get("helper_text"):
                    meta_lines.append(f"helper: {q['helper_text']}")
                meta_lines.append(
                    f"source: {q['source_file']}:{q['source_line']}"
                )
                meta_p = doc.add_paragraph()
                meta_p.add_run("\n".join(meta_lines)).font.size = Pt(9)

    # Gaps
    h = doc.add_heading(f"{section_index}. Gaps and observations", level=1)
    _set_heading_color(h, HEADER_FILL_HEX)
    section_index += 1
    for g in GAPS:
        doc.add_paragraph(g, style="List Bullet")

    # Notes for review
    h = doc.add_heading(f"{section_index}. Notes for review", level=1)
    _set_heading_color(h, HEADER_FILL_HEX)
    doc.add_paragraph(
        "Decisions are recorded in the Excel companion. On the \"Re-ordering Workspace\" sheet, "
        "rows highlight yellow when Suggested New Module differs from the current Module, "
        "green when Karen's Verdict is \"Keep\", and red when it is \"Remove\". The Word "
        "document is read-only reference."
    )

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_PATH)


def main() -> None:
    build_excel()
    build_word()
    print(f"Wrote {XLSX_PATH}")
    print(f"Wrote {DOCX_PATH}")
    print(f"Total questions: {len(QUESTIONS)}")


if __name__ == "__main__":
    main()
